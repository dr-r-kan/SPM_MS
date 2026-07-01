from pathlib import Path
import base64
import math
from html.parser import HTMLParser
import re

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_DOCX = Path(
    "/Users/rohan/Library/CloudStorage/OneDrive-NHS/Microstates/SPM_MS.docx"
)
OUT_DIR = ROOT / "outputs" / "manuscripts"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "SPM_MS_current_pipeline_manuscript.docx"


SIM_METHODS_HTML = ROOT / "outputs/simulations/results/simulation_methods_report.html"
SIM_DESCRIPTIVE_HTML = ROOT / "outputs/simulations/results/simulation_descriptive_results.html"
SIM_BENCHMARK_HTML = ROOT / "outputs/simulations/results/simulation_benchmark_results.html"
SIM_FIG_DIR = ROOT / "outputs/simulations/analysis_plots_python_notebook"
LEMON_DIR = ROOT / "outputs/hierarchical_microstates"
BEHAV_DIR = LEMON_DIR / "behavioural_backfit_analysis"

PRIMARY_SPM_CRITERION = "icl"
PRIMARY_SPM_CRITERION_TEXT = "matrix-space ICL"
PRIMARY_SPM_TABLE_LABEL = "SPM-MS/ICL"
PRIMARY_KM_CRITERION = "silhouette"
PRIMARY_KM_TABLE_LABEL = "K-means/sil."

PAIR_COLUMNS = [
    "rep",
    "K_true",
    "SNR_dB",
    "overlap_prob",
    "overlap_strength",
    "overlap_ms_min",
    "overlap_ms_max",
    "true_template_labels",
    "true_template_indices",
    "montage_type",
    "n_leads",
]

BLOCK_COLUMNS = [
    "rep",
    "K_true",
    "SNR_dB",
    "overlap_prob",
    "overlap_strength",
    "overlap_ms_min",
    "overlap_ms_max",
    "true_template_labels",
    "true_template_indices",
]

PRIMARY_METRICS = [
    ("K_correct", "Exact K recovery", "higher", "rate"),
    ("K_abs_error", "Absolute K error", "lower", "count"),
    ("K_sq_error", "Squared K error", "lower", "count"),
    ("f1_score", "Recovered-state F1", "higher", "rate"),
    ("sensitivity", "Sensitivity", "higher", "rate"),
    ("precision", "Precision", "higher", "rate"),
    ("mean_recovery_padded", "Map recovery, penalized", "higher", "score"),
    ("mean_recovery_matched", "Map recovery, matched", "higher", "score"),
    ("cluster_identity_accuracy", "Template identity accuracy", "higher", "rate"),
    ("backfit_mix_label_top1_accuracy", "Backfit label accuracy", "higher", "rate"),
    ("backfit_mix_label_weight_mae", "Backfit label-weight MAE", "lower", "score"),
    ("backfit_mix_label_pair_accuracy_overlap", "Overlap pair accuracy", "higher", "rate"),
    ("backfit_mix_label_weight_mae_overlap", "Overlap label-weight MAE", "lower", "score"),
    ("runtime_s", "Runtime", "lower", "seconds"),
]


def clean_text(value):
    if value is None:
        return ""
    if pd.isna(value):
        return ""
    return str(value)


def fmt_num(value, nd=3):
    if value is None or pd.isna(value):
        return ""
    return f"{float(value):.{nd}f}"


def fmt_pct(value):
    if value is None or pd.isna(value):
        return ""
    return f"{100 * float(value):.1f}%"


def fmt_p(value):
    if value is None or pd.isna(value):
        return ""
    value = float(value)
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.3f}"


def format_metric_value(value, kind):
    if value is None or not np.isfinite(float(value)):
        return ""
    value = float(value)
    if kind == "rate":
        return f"{100 * value:.1f}%"
    if kind == "seconds":
        return f"{value:.1f}s"
    return f"{value:.3f}" if abs(value) < 10 else f"{value:.2f}"


def format_metric_signed(value, kind):
    if value is None or not np.isfinite(float(value)):
        return ""
    value = float(value)
    sign = "+" if value >= 0 else ""
    if kind == "rate":
        return f"{sign}{100 * value:.1f} pp"
    if kind == "seconds":
        return f"{sign}{value:.1f}s"
    return f"{sign}{value:.3f}" if abs(value) < 10 else f"{sign}{value:.2f}"


def format_metric_p(value, floor=None):
    if value is None or not np.isfinite(float(value)):
        return ""
    value = float(value)
    if floor is not None and np.isfinite(float(floor)) and value <= float(floor) + 1e-15:
        return f"<{float(floor):.1e}"
    if value < 0.001:
        return f"{value:.1e}"
    return f"{value:.3f}"


def clean_token(value):
    text = "" if value is None or pd.isna(value) else str(value)
    return re.sub(r"\s+", " ", text.lower().strip().replace("_", " ").replace("-", " "))


def criterion_short_label(criterion):
    names = {
        "calinski_harabasz_score": "Calinski-Harabasz",
        "covariance_elbow": "covariance elbow",
        "elbow_sil_combined": "elbow + silhouette",
        "free_energy": "free energy",
        "log_likelihood": "LL",
        "bic": "BIC",
        "icl": "ICL",
        "free_energy_elbow": "free-energy elbow",
        "free_energy_covariance": "FE + covariance",
        "gev": "GEV",
        "silhouette": "silhouette",
    }
    return names.get(str(criterion), str(criterion).replace("_", " "))


def add_omml_like_run(paragraph, text, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_in):
    width_twips = int(width_in * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table):
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_cell_margins(table, top=45, start=55, bottom=45, end=55):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m_name, m_value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        node = margins.find(qn(f"w:{m_name}"))
        if node is None:
            node = OxmlElement(f"w:{m_name}")
            margins.append(node)
        node.set(qn("w:w"), str(m_value))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    title = styles["Title"]
    title.font.name = "Calibri Light"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri Light")
    title.font.size = Pt(20)
    title.font.color.rgb = RGBColor(31, 78, 121)
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before in [
        ("Heading 1", 15, RGBColor(31, 78, 121), 14),
        ("Heading 2", 12.5, RGBColor(46, 87, 112), 10),
        ("Heading 3", 11.2, RGBColor(64, 64, 64), 8),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(4)

    if "Caption" not in styles:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    cap = styles["Caption"]
    cap.font.name = "Calibri"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(80, 80, 80)
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)

    if "TableText" not in styles:
        styles.add_style("TableText", WD_STYLE_TYPE.PARAGRAPH)
    tab = styles["TableText"]
    tab.font.name = "Calibri"
    tab._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    tab.font.size = Pt(8.2)
    tab.paragraph_format.space_after = Pt(0)
    tab.paragraph_format.line_spacing = 1.0


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    if align:
        p.alignment = align
    return p


def add_bold_lead(doc, lead, rest):
    p = doc.add_paragraph()
    p.add_run(lead).bold = True
    p.add_run(rest)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    set_cell_margins(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = clean_text(header)
        set_cell_shading(hdr[idx], "D9EAF7")
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(hdr[idx], widths[idx])
        for p in hdr[idx].paragraphs:
            p.style = "TableText"
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = clean_text(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                set_cell_width(cells[idx], widths[idx])
            for p in cells[idx].paragraphs:
                p.style = "TableText"
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    doc.add_paragraph(text, style="Caption")


def add_picture(doc, path, width_in=6.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))


def read_bibliography():
    if not SOURCE_DOCX.exists():
        return []
    src = Document(str(SOURCE_DOCX))
    refs = []
    for para in src.paragraphs:
        if para.style.name == "Bibliography" and para.text.strip():
            refs.append(re.sub(r"\s+", " ", para.text.strip()))
    return refs


class ReportImageParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.images = []

    def handle_starttag(self, tag, attrs):
        if tag == "img":
            self.images.append(dict(attrs))


def slugify(value):
    return re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")


def extract_report_images(report_path, wanted_alts, prefix):
    out_dir = ASSET_DIR / "report_images"
    out_dir.mkdir(parents=True, exist_ok=True)
    parser = ReportImageParser()
    parser.feed(report_path.read_text(errors="ignore"))
    wanted = set(wanted_alts)
    found = {}
    for image in parser.images:
        alt = image.get("alt", "")
        src = image.get("src", "")
        if alt not in wanted or not src.startswith("data:image/png;base64,"):
            continue
        path = out_dir / f"{prefix}_{slugify(alt)}.png"
        path.write_bytes(base64.b64decode(src.split(",", 1)[1]))
        found[alt] = path
    return found


def build_lemon_topography_panel():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    out = ASSET_DIR / "lemon_topography_panel.png"
    source_paths = [
        ("Global", LEMON_DIR / "global/global_centers.png"),
        ("Eyes closed", LEMON_DIR / "conditions/EC/condition_centers.png"),
        ("Eyes open", LEMON_DIR / "conditions/EO/condition_centers.png"),
    ]
    images = []
    for label, path in source_paths:
        img = Image.open(path).convert("RGB")
        images.append((label, img))

    canvas_w = 1700
    pad = 28
    label_h = 46
    font = ImageFont.load_default()

    def resize_to_width(img, width):
        ratio = width / img.width
        return img.resize((width, int(img.height * ratio)), Image.Resampling.LANCZOS)

    global_img = resize_to_width(images[0][1], canvas_w - 2 * pad)
    child_w = (canvas_w - 3 * pad) // 2
    ec_img = resize_to_width(images[1][1], child_w)
    eo_img = resize_to_width(images[2][1], child_w)
    canvas_h = (
        pad
        + label_h
        + global_img.height
        + pad
        + label_h
        + max(ec_img.height, eo_img.height)
        + pad
    )
    canvas = Image.new("RGB", (canvas_w, canvas_h), "white")
    draw = ImageDraw.Draw(canvas)

    y = pad
    draw.text((pad, y), "Global", fill=(31, 78, 121), font=font)
    y += label_h
    canvas.paste(global_img, (pad, y))
    y += global_img.height + pad
    draw.text((pad, y), "Eyes closed", fill=(31, 78, 121), font=font)
    draw.text((2 * pad + child_w, y), "Eyes open", fill=(31, 78, 121), font=font)
    y += label_h
    canvas.paste(ec_img, (pad, y))
    canvas.paste(eo_img, (2 * pad + child_w, y))
    canvas.save(out, dpi=(220, 220))
    return out


def build_pipeline_demonstrator_panel():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    source = ASSET_DIR / "pipeline_demo_simulated_eeg.csv"
    out = ASSET_DIR / "pipeline_demonstrator_panel.png"
    eeg = pd.read_csv(source)
    eeg = eeg[eeg["time_s"] <= 4.0].copy()

    def font(size, bold=False):
        paths = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        ]
        for path in paths:
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                pass
        return ImageFont.load_default()

    def text_size(draw, text, fnt):
        box = draw.textbbox((0, 0), text, font=fnt)
        return box[2] - box[0], box[3] - box[1]

    def wrapped_lines(draw, text, fnt, width):
        lines = []
        for part in text.split("\n"):
            cur = ""
            for word in part.split():
                trial = f"{cur} {word}".strip()
                if cur and text_size(draw, trial, fnt)[0] > width:
                    lines.append(cur)
                    cur = word
                else:
                    cur = trial
            lines.append(cur)
        return lines

    def draw_centered_text(draw, box, text, fnt, fill):
        x0, y0, x1, y1 = box
        lines = wrapped_lines(draw, text, fnt, x1 - x0 - 36)
        line_h = text_size(draw, "Ag", fnt)[1] + 8
        y = y0 + ((y1 - y0) - line_h * len(lines)) / 2
        for line in lines:
            w, _ = text_size(draw, line, fnt)
            draw.text((x0 + ((x1 - x0) - w) / 2, y), line, font=fnt, fill=fill)
            y += line_h

    def arrow(draw, x0, y0, x1, y1, fill):
        draw.line((x0, y0, x1, y1), fill=fill, width=5)
        head = 16
        draw.polygon([(x1, y1), (x1 - head, y1 - head), (x1 + head, y1 - head)], fill=fill)

    W, H = 2200, 1450
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    title_f, head_f, body_f, small_f = font(38, True), font(26, True), font(21), font(18)
    ink, muted, grid = (35, 35, 35), (95, 95, 95), (220, 220, 220)
    blue, orange = (78, 121, 167), (242, 142, 43)
    state_colors = {1: (78, 121, 167), 2: (242, 142, 43), 3: (89, 161, 79), 4: (176, 122, 161)}
    state_labels = {1: "A", 2: "B", 3: "E", 4: "G"}

    draw.text((70, 42), "Pipeline demonstrator: simulated EEG to microstate outputs", font=title_f, fill=ink)
    draw.text((80, 120), "A. Simulated EEG excerpt", font=head_f, fill=ink)

    x0, y0, plot_w, plot_h = 115, 190, 1280, 555
    channels = ["Fz", "Cz", "Pz", "O1"]
    time = eeg["time_s"].astype(float).to_numpy()
    t_min, t_max = float(time.min()), float(time.max())
    amp = max(float(eeg[channels].abs().quantile(0.98).max()), 1.0)
    row_h = plot_h / len(channels)

    for tick in range(0, 5):
        x = x0 + int((tick - t_min) / (t_max - t_min) * plot_w)
        draw.line((x, y0, x, y0 + plot_h), fill=grid, width=2)
        draw.text((x - 8, y0 + plot_h + 12), str(tick), font=small_f, fill=muted)

    for i, ch in enumerate(channels):
        base_y = y0 + row_h * (i + 0.5)
        draw.line((x0, base_y, x0 + plot_w, base_y), fill=(238, 238, 238), width=2)
        draw.text((34, base_y - 13), ch, font=body_f, fill=ink)
        values = eeg[ch].astype(float).to_numpy()
        pts = []
        for t, value in zip(time, values):
            x = x0 + int((t - t_min) / (t_max - t_min) * plot_w)
            y = base_y - (float(value) / amp) * row_h * 0.34
            pts.append((x, int(y)))
        draw.line(pts, fill=blue if i % 2 == 0 else orange, width=3)

    draw.rectangle((x0, y0, x0 + plot_w, y0 + plot_h), outline=(90, 90, 90), width=3)
    draw.text((x0 + plot_w // 2 - 55, y0 + plot_h + 48), "Time (s)", font=body_f, fill=ink)
    draw.text((x0, y0 + plot_h + 85), "Example: K=true 4, SNR=-1 dB, 50% boundary-overlap probability, 250 Hz", font=small_f, fill=muted)

    state_y0, state_h = y0 + plot_h + 130, 78
    states = eeg["true_state"].astype(int).to_numpy()
    start = 0
    for idx in range(1, len(states) + 1):
        if idx == len(states) or states[idx] != states[start]:
            x_start = x0 + int((time[start] - t_min) / (t_max - t_min) * plot_w)
            x_end = x0 + int((time[idx - 1] - t_min) / (t_max - t_min) * plot_w)
            state = int(states[start])
            draw.rectangle((x_start, state_y0, max(x_start + 2, x_end), state_y0 + state_h), fill=state_colors[state])
            if x_end - x_start > 34:
                draw_centered_text(draw, (x_start, state_y0, x_end, state_y0 + state_h), state_labels[state], small_f, "white")
            start = idx
    draw.rectangle((x0, state_y0, x0 + plot_w, state_y0 + state_h), outline=(90, 90, 90), width=3)
    draw.text((34, state_y0 + 20), "True\nstate", font=small_f, fill=ink)

    flow_x0, flow_y0 = 1510, 160
    draw.text((flow_x0, 120), "B. Pipeline flow", font=head_f, fill=ink)
    boxes = [
        "Generate simulated EEG\nknown K, maps, SNR,\noverlap labels",
        "Preprocess\nmontage projection,\nGFP peaks",
        "Fit candidate K\nSPM-MS VB mixture\nand K-means",
        "Select and align\ncriterion curves,\ntemplate labels",
        "Backfit timecourse\nstate labels,\nstate weights",
        "Report outputs\nsquared K error, F1,\nruntime, LEMON hierarchy",
    ]
    box_w, box_h, gap = 560, 132, 42
    for i, label in enumerate(boxes):
        bx0 = flow_x0
        by0 = flow_y0 + i * (box_h + gap)
        fill = (232, 239, 247) if i % 2 == 0 else (253, 239, 219)
        draw.rounded_rectangle((bx0, by0, bx0 + box_w, by0 + box_h), radius=18, fill=fill, outline=(110, 110, 110), width=3)
        draw_centered_text(draw, (bx0, by0, bx0 + box_w, by0 + box_h), label, body_f, ink)
        if i < len(boxes) - 1:
            arrow(draw, bx0 + box_w / 2, by0 + box_h + 4, bx0 + box_w / 2, by0 + box_h + gap - 8, (85, 85, 85))

    draw.text((80, H - 80), "The trace is regenerated with the same simulation generator used for the benchmark; downstream panels summarise the full experimental run.", font=small_f, fill=muted)
    img.save(out, dpi=(220, 220))
    return out


def display_method_criterion(method, criterion):
    method_name = "K means" if method == "kmeans_koenig" else "SPM-MS"
    names = {
        "calinski_harabasz_score": "Calinski-Harabasz",
        "covariance_elbow": "Covariance elbow",
        "elbow_sil_combined": "Elbow + silhouette",
        "free_energy": "Free energy",
        "log_likelihood": "LL",
        "bic": "BIC",
        "icl": "ICL",
        "free_energy_elbow": "Free-energy elbow",
        "free_energy_covariance": "FE + covariance",
        "gev": "GEV",
    }
    criterion_name = names.get(criterion, criterion.replace("_", " ").title())
    return f"{method_name}: {criterion_name}"


def build_simulation_diagnostic_plots(comparison, spm_criterion=PRIMARY_SPM_CRITERION):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    comp = comparison.copy()
    comp["display"] = [
        display_method_criterion(m, c) for m, c in zip(comp["method"], comp["criterion"])
    ]
    comp["K_est_minus_true"] = comp["K_estimated"] - comp["K_true"]
    order = [
        ("kmeans_koenig", "silhouette"),
        ("spm_vb", "free_energy"),
        ("spm_vb", "log_likelihood"),
        ("spm_vb", "bic"),
        ("spm_vb", "icl"),
        ("spm_vb", "calinski_harabasz_score"),
        ("spm_vb", "silhouette"),
        ("spm_vb", "covariance"),
        ("spm_vb", "gev"),
        ("spm_vb", "covariance_elbow"),
        ("spm_vb", "free_energy_elbow"),
        ("spm_vb", "elbow_sil_combined"),
        ("spm_vb", "free_energy_covariance"),
    ]
    label_order = [display_method_criterion(m, c) for m, c in order]
    summary = (
        comp.groupby("display", observed=True)
        .agg(
            mean_k=("K_estimated", "mean"),
            se_k=("K_estimated", lambda x: x.std(ddof=1) / math.sqrt(len(x))),
            mean_gap=("K_est_minus_true", "mean"),
            se_gap=("K_est_minus_true", lambda x: x.std(ddof=1) / math.sqrt(len(x))),
        )
        .reindex(label_order)
    )

    def font(size, bold=False):
        paths = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        ]
        for path in paths:
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                pass
        return ImageFont.load_default()

    W, H = 2200, 1180
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    f_title, f_axis, f_small = font(34, True), font(23), font(20)
    left_label = 35
    panel1_x, panel2_x = 650, 1420
    panel_w = 620
    top, row_h = 130, 88
    blue, orange, grid, text = (78, 121, 167), (242, 142, 43), (220, 220, 220), (35, 35, 35)
    draw.text((panel1_x, 35), "Selected K by criterion", fill=text, font=f_title)
    draw.text((panel2_x, 35), "Direction of K error", fill=text, font=f_title)
    true_mean = comp["K_true"].mean()
    k_min, k_max = 0, 10
    gap_min, gap_max = -4, 5

    def xscale(value, xmin, xmax, x0):
        return x0 + int((float(value) - xmin) / (xmax - xmin) * panel_w)

    for tick in range(k_min, k_max + 1, 2):
        x = xscale(tick, k_min, k_max, panel1_x)
        draw.line((x, top - 20, x, top + row_h * len(summary)), fill=grid)
        draw.text((x - 10, top + row_h * len(summary) + 12), str(tick), fill=text, font=f_small)
    for tick in range(gap_min, gap_max + 1):
        x = xscale(tick, gap_min, gap_max, panel2_x)
        draw.line((x, top - 20, x, top + row_h * len(summary)), fill=grid)
        draw.text((x - 12, top + row_h * len(summary) + 12), str(tick), fill=text, font=f_small)
    ref_x = xscale(true_mean, k_min, k_max, panel1_x)
    draw.line((ref_x, top - 28, ref_x, top + row_h * len(summary)), fill=(60, 60, 60), width=3)
    zero_x = xscale(0, gap_min, gap_max, panel2_x)
    draw.line((zero_x, top - 28, zero_x, top + row_h * len(summary)), fill=(60, 60, 60), width=3)

    for i, (label, row) in enumerate(summary.iterrows()):
        y = top + i * row_h + 18
        color = blue if label.startswith("K means") else orange
        draw.text((left_label, y + 4), label, fill=text, font=f_axis)
        for x0, xmin, xmax, mean_col, se_col in [
            (panel1_x, k_min, k_max, "mean_k", "se_k"),
            (panel2_x, gap_min, gap_max, "mean_gap", "se_gap"),
        ]:
            base_x = xscale(0, xmin, xmax, x0)
            value_x = xscale(row[mean_col], xmin, xmax, x0)
            x1, x2 = sorted([base_x, value_x])
            draw.rectangle((x1, y, x2, y + 34), fill=color)
            ci = 1.96 * row[se_col]
            err_l = xscale(row[mean_col] - ci, xmin, xmax, x0)
            err_r = xscale(row[mean_col] + ci, xmin, xmax, x0)
            mid_y = y + 17
            draw.line((err_l, mid_y, err_r, mid_y), fill=text, width=3)
            draw.line((err_l, mid_y - 8, err_l, mid_y + 8), fill=text, width=3)
            draw.line((err_r, mid_y - 8, err_r, mid_y + 8), fill=text, width=3)
            label_x = max(x1, x2) + 8
            if value_x < base_x:
                label_x = value_x - 58
            draw.text((label_x, y + 3), f"{row[mean_col]:.2f}", fill=text, font=f_small)
    draw.text((panel1_x, H - 55), "Mean selected K", fill=text, font=f_axis)
    draw.text((panel2_x, H - 55), "Mean signed K difference (estimated - true)", fill=text, font=f_axis)
    criterion_path = ASSET_DIR / "simulation_criterion_selected_k_and_gap.png"
    img.save(criterion_path, dpi=(220, 220))

    best = comp[
        ((comp["method"] == "kmeans_koenig") & (comp["criterion"] == "silhouette"))
        | ((comp["method"] == "spm_vb") & (comp["criterion"] == spm_criterion))
    ].copy()
    spm_series_label = f"SPM-MS / {criterion_short_label(spm_criterion)}"
    best["method_label"] = best["method"].map(
        {"kmeans_koenig": "K means / silhouette", "spm_vb": spm_series_label}
    )
    f1 = (
        best.groupby(["method_label", "K_true"], observed=True)["f1_score"]
        .agg(["mean", "std", "count"])
        .reset_index()
    )
    f1["ci"] = 1.96 * f1["std"] / f1["count"].map(math.sqrt)
    W, H = 1500, 900
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    left, right, top, bottom = 150, 70, 105, 120
    plot_w, plot_h = W - left - right, H - top - bottom
    draw.text((left, 35), "Recoverable-structure criterion: F1 by true K", fill=text, font=f_title)

    def x_k(k):
        return left + int((int(k) - 4) / 3 * plot_w)

    def y_f1(v):
        v = max(0.5, min(1.02, float(v)))
        return top + int((1.02 - v) / (1.02 - 0.5) * plot_h)

    for k in [4, 5, 6, 7]:
        x = x_k(k)
        draw.line((x, top, x, top + plot_h), fill=grid)
        draw.text((x - 8, top + plot_h + 22), str(k), fill=text, font=f_axis)
    for yval in [0.5, 0.6, 0.7, 0.8, 0.9, 1.0]:
        y = y_f1(yval)
        draw.line((left, y, left + plot_w, y), fill=grid)
        draw.text((left - 58, y - 12), f"{yval:.1f}", fill=text, font=f_axis)
    draw.line((left, top, left, top + plot_h), fill=text, width=3)
    draw.line((left, top + plot_h, left + plot_w, top + plot_h), fill=text, width=3)
    series_colors = {"K means / silhouette": blue, spm_series_label: orange}
    legend_y = top + plot_h - 90
    for label, group in f1.groupby("method_label", sort=False):
        group = group.sort_values("K_true")
        color = series_colors[label]
        pts = [(x_k(r["K_true"]), y_f1(r["mean"])) for _, r in group.iterrows()]
        upper = [(x_k(r["K_true"]), y_f1(r["mean"] + r["ci"])) for _, r in group.iterrows()]
        lower = [(x_k(r["K_true"]), y_f1(r["mean"] - r["ci"])) for _, r in group.iterrows()]
        overlay = Image.new("RGBA", (W, H), (255, 255, 255, 0))
        ImageDraw.Draw(overlay).polygon(upper + list(reversed(lower)), fill=(*color, 48))
        img = Image.alpha_composite(img.convert("RGBA"), overlay).convert("RGB")
        draw = ImageDraw.Draw(img)
        draw.line(pts, fill=color, width=5)
        for _, r in group.iterrows():
            x, y = x_k(r["K_true"]), y_f1(r["mean"])
            y0, y1 = y_f1(r["mean"] + r["ci"]), y_f1(r["mean"] - r["ci"])
            draw.line((x, y0, x, y1), fill=color, width=3)
            draw.line((x - 8, y0, x + 8, y0), fill=color, width=3)
            draw.line((x - 8, y1, x + 8, y1), fill=color, width=3)
            draw.ellipse((x - 11, y - 11, x + 11, y + 11), fill=color, outline="white", width=3)
        draw.rectangle((left + 25, legend_y, left + 65, legend_y + 18), fill=color)
        draw.text((left + 78, legend_y - 5), label, fill=text, font=f_axis)
        legend_y += 40
    draw.text((left + plot_w // 2 - 45, H - 58), "True K", fill=text, font=f_axis)
    draw.text((25, top + plot_h // 2 - 20), "F1", fill=text, font=f_axis)
    f1_path = ASSET_DIR / "simulation_best_criterion_f1_by_true_k.png"
    img.save(f1_path, dpi=(220, 220))
    return criterion_path, f1_path


def build_selected_criterion_k_confusion(comparison, spm_criterion=PRIMARY_SPM_CRITERION):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    out = ASSET_DIR / "simulation_selected_criterion_k_confusion.png"

    def font(size, bold=False):
        paths = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        ]
        for path in paths:
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                pass
        return ImageFont.load_default()

    def blue(value):
        value = max(0.0, min(1.0, float(value)))
        lo, hi = (245, 249, 252), (8, 81, 156)
        return tuple(int(lo[i] + value * (hi[i] - lo[i])) for i in range(3))

    def text_size(draw, text, fnt):
        box = draw.textbbox((0, 0), str(text), font=fnt)
        return box[2] - box[0], box[3] - box[1]

    def centered(draw, box, text, fnt, fill):
        x0, y0, x1, y1 = box
        w, h = text_size(draw, text, fnt)
        draw.text((x0 + (x1 - x0 - w) / 2, y0 + (y1 - y0 - h) / 2), str(text), font=fnt, fill=fill)

    W, H = 1900, 1600
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    title_f, head_f, axis_f, small_f = font(34, True), font(23, True), font(19), font(16)
    text, grid = (35, 35, 35), (205, 205, 205)
    draw.text((70, 35), "K-selection confusion: selected criteria only", font=title_f, fill=text)
    draw.text(
        (70, 85),
        f"Rows are true K; columns are estimated K. SPM-MS is filtered to {criterion_short_label(spm_criterion)}, not averaged across criteria.",
        font=axis_f,
        fill=(80, 80, 80),
    )

    montages = [("full", "full (71 ch)"), ("10-20-20", "10-20-20 (19 ch)"), ("10-20-12", "10-20-12 (12 ch)")]
    methods = [
        ("kmeans_koenig", "silhouette", "K-means / silhouette"),
        ("spm_vb", spm_criterion, f"SPM-MS / {criterion_short_label(spm_criterion)}"),
    ]
    panel_w, panel_h = 720, 330
    x_starts = [150, 1030]
    y_starts = [180, 620, 1060]
    true_vals = [4, 5, 6, 7]

    for row, (montage, montage_label) in enumerate(montages):
        for col, (method, criterion, method_label) in enumerate(methods):
            sub = comparison[
                (comparison["montage_type"].astype(str) == montage)
                & (comparison["method"].astype(str) == method)
                & (comparison["criterion"].astype(str) == criterion)
            ].copy()
            if sub.empty:
                continue
            est_vals = sorted(int(x) for x in pd.to_numeric(sub["K_estimated"], errors="coerce").dropna().unique())
            x0, y0 = x_starts[col], y_starts[row]
            draw.text((x0, y0 - 38), f"{montage_label} | {method_label}", font=head_f, fill=text)
            left, top = x0 + 70, y0 + 15
            heat_w, heat_h = panel_w - 120, panel_h - 95
            cell_w, cell_h = heat_w / len(est_vals), heat_h / len(true_vals)
            tab = pd.crosstab(
                pd.to_numeric(sub["K_true"], errors="coerce").astype(int),
                pd.to_numeric(sub["K_estimated"], errors="coerce").astype(int),
                normalize="index",
            ).reindex(index=true_vals, columns=est_vals, fill_value=0)
            for r, k_true in enumerate(true_vals):
                y1 = top + r * cell_h
                draw.text((left - 38, y1 + cell_h / 2 - 10), str(k_true), font=axis_f, fill=text)
                for c, k_est in enumerate(est_vals):
                    val = float(tab.loc[k_true, k_est])
                    x1 = left + c * cell_w
                    box = (x1, y1, x1 + cell_w, y1 + cell_h)
                    draw.rectangle(box, fill=blue(val), outline=grid)
                    if val >= 0.05:
                        fill = "white" if val > 0.55 else text
                        centered(draw, box, f"{100 * val:.0f}%", small_f, fill)
            for c, k_est in enumerate(est_vals):
                centered(draw, (left + c * cell_w, top + heat_h + 8, left + (c + 1) * cell_w, top + heat_h + 36), str(k_est), axis_f, text)
            draw.text((left + heat_w / 2 - 55, top + heat_h + 46), "Estimated K", font=axis_f, fill=text)
            draw.text((x0, top + heat_h / 2 - 12), "True K", font=axis_f, fill=text)
    img.save(out, dpi=(220, 220))
    return out


def bootstrap_ci(values, n_boot, rng):
    values = np.asarray(values, dtype=float)
    values = values[np.isfinite(values)]
    if values.size == 0:
        return np.nan, np.nan
    samples = values[rng.integers(0, values.size, size=(n_boot, values.size))].mean(axis=1)
    return tuple(np.percentile(samples, [2.5, 97.5]))


def sign_flip_p_value(values, n_perm, rng):
    values = np.asarray(values, dtype=float)
    values = values[np.isfinite(values)]
    if values.size < 2:
        return np.nan, np.nan
    observed = values.mean()
    exceed = 0
    done = 0
    chunk = 5000
    while done < n_perm:
        size = min(chunk, n_perm - done)
        signs = rng.choice(np.array([-1.0, 1.0]), size=(size, values.size))
        exceed += int(np.sum(signs.dot(values) / values.size >= observed - 1e-12))
        done += size
    floor = 1.0 / (n_perm + 1)
    return (exceed + 1.0) / (n_perm + 1), floor


def add_holm_p(df):
    df = df.copy()
    pvals = pd.to_numeric(df["p_value"], errors="coerce").to_numpy(float)
    adjusted = np.full(len(df), np.nan)
    finite = np.flatnonzero(np.isfinite(pvals))
    order = finite[np.argsort(pvals[finite])]
    running = 0.0
    total = len(order)
    for rank, idx in enumerate(order):
        running = max(running, (total - rank) * pvals[idx])
        adjusted[idx] = min(running, 1.0)
    df["p_holm"] = adjusted
    return df


def paired_primary_benchmark(comparison, spm_criterion=PRIMARY_SPM_CRITERION):
    df = comparison.copy()
    df["method_clean"] = df["method"].map(clean_token)
    df["criterion_clean"] = df["criterion"].map(clean_token)
    k_error = pd.to_numeric(df["K_estimated"], errors="coerce") - pd.to_numeric(df["K_true"], errors="coerce")
    df["K_abs_error"] = k_error.abs()
    df["K_sq_error"] = k_error.pow(2)

    spm = df[(df["method_clean"].str.contains("spm")) & (df["criterion_clean"].eq(clean_token(spm_criterion)))]
    km = df[
        (df["method_clean"].str.contains("kmeans") | df["method_clean"].str.contains("k means"))
        & (df["criterion_clean"].eq(clean_token(PRIMARY_KM_CRITERION)))
    ]
    if spm.empty or km.empty:
        raise ValueError(f"Missing primary comparison rows for {spm_criterion} versus {PRIMARY_KM_CRITERION}.")

    metric_cols = sorted({m[0] for m in PRIMARY_METRICS if m[0] in df.columns} | {"K_estimated"})
    spm = spm[PAIR_COLUMNS + metric_cols].set_index(PAIR_COLUMNS).sort_index()
    km = km[PAIR_COLUMNS + metric_cols].set_index(PAIR_COLUMNS).sort_index()
    common = spm.index.intersection(km.index)
    pairs = spm.loc[common].join(km.loc[common], lsuffix="_spm", rsuffix="_km").reset_index()
    if len(pairs) != len(spm) or len(pairs) != len(km):
        raise ValueError(f"Only {len(pairs)} paired rows found from {len(spm)} SPM-MS and {len(km)} K-means rows.")

    rows = []
    rng = np.random.default_rng(0)
    montage_groups = []
    for montage, group in pairs.groupby("montage_type", sort=False):
        n_leads = pd.to_numeric(group["n_leads"], errors="coerce").dropna()
        montage_groups.append((int(n_leads.iloc[0]) if not n_leads.empty else -1, montage, group))

    for _, montage, group in sorted(montage_groups, key=lambda item: (-item[0], str(item[1]))):
        n_leads = int(pd.to_numeric(group["n_leads"], errors="coerce").dropna().iloc[0])
        montage_label = f"{montage} ({n_leads} ch)"
        for col, label, direction, kind in PRIMARY_METRICS:
            spm_col, km_col = f"{col}_spm", f"{col}_km"
            if spm_col not in group or km_col not in group:
                continue
            spm_values = pd.to_numeric(group[spm_col], errors="coerce")
            km_values = pd.to_numeric(group[km_col], errors="coerce")
            keep = spm_values.notna() & km_values.notna()
            if not keep.any():
                continue

            improvement = spm_values - km_values if direction == "higher" else km_values - spm_values
            tmp = group.loc[keep, BLOCK_COLUMNS].assign(_improvement=improvement[keep])
            block_improvement = tmp.groupby(BLOCK_COLUMNS, dropna=False)["_improvement"].mean().to_numpy(float)
            seed_improvement = (
                group.loc[keep, ["rep"]]
                .assign(_improvement=improvement[keep])
                .groupby(["rep"], dropna=False)["_improvement"]
                .mean()
                .to_numpy(float)
            )
            lo, hi = bootstrap_ci(block_improvement, 2000, rng)
            p_value, p_floor = sign_flip_p_value(block_improvement, 20000, rng)
            wins = int(np.sum(block_improvement > 1e-12))
            ties = int(np.sum(np.abs(block_improvement) <= 1e-12))
            losses = int(np.sum(block_improvement < -1e-12))
            rows.append(
                {
                    "Montage": montage_label,
                    "Metric": label,
                    "Direction": "higher is better" if direction == "higher" else "lower is better",
                    "Rows": f"{int(keep.sum()):,}",
                    "Blocks": f"{len(block_improvement):,}",
                    "K means": format_metric_value(km_values[keep].mean(), kind),
                    "SPM-MS": format_metric_value(spm_values[keep].mean(), kind),
                    "Improvement": format_metric_signed(np.mean(block_improvement), kind),
                    "95% CI": f"{format_metric_signed(lo, kind)} to {format_metric_signed(hi, kind)}",
                    "Permutation p": format_metric_p(p_value, p_floor),
                    "Holm p": "",
                    "dz": "",
                    "Block W/T/L": f"{wins}/{ties}/{losses}",
                    "Seed W/T/L": (
                        f"{int(np.sum(seed_improvement > 1e-12))}/"
                        f"{int(np.sum(np.abs(seed_improvement) <= 1e-12))}/"
                        f"{int(np.sum(seed_improvement < -1e-12))}"
                    ),
                    "_p_value": p_value,
                    "_kind": kind,
                }
            )

    out = add_holm_p(pd.DataFrame(rows).rename(columns={"_p_value": "p_value"}))
    out["Holm p"] = [format_metric_p(v) for v in out["p_holm"]]
    return out.drop(columns=["p_value", "p_holm", "_kind"], errors="ignore")


def spm_criterion_summary(comparison):
    spm = comparison[comparison["method"].astype(str).eq("spm_vb")].copy()
    spm["K_est_minus_true"] = pd.to_numeric(spm["K_estimated"], errors="coerce") - pd.to_numeric(
        spm["K_true"], errors="coerce"
    )
    spm["K_abs_error"] = spm["K_est_minus_true"].abs()
    spm["K_sq_error"] = spm["K_est_minus_true"].pow(2)
    spm["under_selected"] = spm["K_est_minus_true"] < 0
    spm["over_selected"] = spm["K_est_minus_true"] > 0
    spm["missed_states"] = (pd.to_numeric(spm["K_true"], errors="coerce") - pd.to_numeric(spm["n_matched"], errors="coerce")).clip(lower=0)
    spm["extra_states"] = (pd.to_numeric(spm["K_estimated"], errors="coerce") - pd.to_numeric(spm["n_matched"], errors="coerce")).clip(lower=0)
    return (
        spm.groupby("criterion", observed=True)
        .agg(
            n=("K_correct", "size"),
            exact_K=("K_correct", "mean"),
            mean_K=("K_estimated", "mean"),
            signed_K=("K_est_minus_true", "mean"),
            abs_K_error=("K_abs_error", "mean"),
            sq_K_error=("K_sq_error", "mean"),
            under_selected=("under_selected", "mean"),
            over_selected=("over_selected", "mean"),
            missed_states=("missed_states", "mean"),
            extra_states=("extra_states", "mean"),
            sensitivity=("sensitivity", "mean"),
            precision=("precision", "mean"),
            F1=("f1_score", "mean"),
            padded_recovery=("mean_recovery_padded", "mean"),
        )
        .sort_values(["F1", "padded_recovery", "sq_K_error"], ascending=[False, False, True])
        .reset_index()
    )


def load_stats():
    methods_tables = pd.read_html(SIM_METHODS_HTML)
    bench_tables = pd.read_html(SIM_BENCHMARK_HTML)
    comparison = pd.read_csv(ROOT / "outputs/simulations/results/comparison_results.csv")
    stats = {
        "design_table": methods_tables[0],
        "montage_table": methods_tables[1],
        "method_table": methods_tables[2],
        "criterion_table": methods_tables[3],
        "outcome_table": methods_tables[4],
        "backfit_table": methods_tables[5],
        "benchmark_checks": bench_tables[0],
        "benchmark": paired_primary_benchmark(comparison, PRIMARY_SPM_CRITERION),
        "criterion_benchmark": bench_tables[2],
        "comparison": comparison,
        "spm_criterion_summary": spm_criterion_summary(comparison),
    }

    manifest = pd.read_csv(LEMON_DIR / "normalised_input_manifest.csv")
    valid = manifest["condition"].isin(["EC", "EO"])
    clean_manifest = manifest[valid].copy()
    common_channels = pd.read_csv(LEMON_DIR / "common_channels.csv")
    fit = pd.read_csv(LEMON_DIR / "hierarchical_fit_summary.csv")
    model = pd.read_csv(LEMON_DIR / "global/global_model_comparison.csv")
    records = pd.read_csv(LEMON_DIR / "participant_condition_record_backfit_summary.csv")
    states = pd.read_csv(LEMON_DIR / "participant_condition_state_backfit_metrics.csv")

    hard_records = records[
        (records["backfit_method"] == "hard")
        & (records["backfit_available"] == 1)
        & (records["condition"].isin(["EC", "EO"]))
    ].copy()
    hard_states = states[
        (states["backfit_method"] == "hard")
        & (states["backfit_available"] == 1)
        & (states["condition"].isin(["EC", "EO"]))
        & (states["occupancy"].notna())
    ].copy()

    stats.update(
        {
            "manifest": manifest,
            "clean_manifest": clean_manifest,
            "common_channels": common_channels,
            "fit": fit,
            "global_model": model,
            "hard_records": hard_records,
            "hard_states": hard_states,
            "pca": pd.read_csv(BEHAV_DIR / "psychometric_pca_summary.csv"),
            "pc_tests": pd.read_csv(BEHAV_DIR / "method_hard/backfit_vs_psychometric_pcs.csv"),
            "demographics": pd.read_csv(BEHAV_DIR / "method_hard/demographic_effects.csv"),
            "focused": pd.read_csv(
                BEHAV_DIR / "method_hard/focused_interoception_dissociation_microstate_stats.csv"
            ),
            "state_e": pd.read_csv(
                BEHAV_DIR / "tas_stai_yfas_state_e_plots/tas_stai_yfas_state_e_correlations.csv"
            ),
        }
    )
    return stats


def metric_row(benchmark, montage, metric):
    row = benchmark[(benchmark["Montage"] == montage) & (benchmark["Metric"] == metric)]
    if row.empty:
        raise ValueError(f"Missing benchmark row: {montage}, {metric}")
    return row.iloc[0]


def add_simulation_results(doc, stats):
    benchmark = stats["benchmark"]
    comparison = stats["comparison"]
    criterion_summary = stats["spm_criterion_summary"]
    criterion_winner = criterion_summary.iloc[0]
    selected_criterion = criterion_summary[criterion_summary["criterion"].eq(PRIMARY_SPM_CRITERION)].iloc[0]
    exact_winner = criterion_summary.sort_values(["exact_K", "F1"], ascending=[False, False]).iloc[0]
    montages = ["full (71 ch)", "10-20-20 (19 ch)", "10-20-12 (12 ch)"]
    criterion_plot, f1_plot = build_simulation_diagnostic_plots(comparison, PRIMARY_SPM_CRITERION)
    confusion_plot = build_selected_criterion_k_confusion(comparison, PRIMARY_SPM_CRITERION)

    add_heading(doc, "Ground-truth simulation recovery", 2)
    add_heading(doc, "Methods report: design and outputs", 3)
    add_paragraph(
        doc,
        "The simulation report contains 560 generated EEG conditions, each evaluated in "
        "three montage projections, giving 1,680 matched generated EEG blocks per method. "
        "Each block was 300 s at 250 Hz, with true K in {4, 5, 6, 7}, SNR in {-6, -3, -1, 0, 1, 3, 6} dB, "
        "and overlap probability 0 or 0.5. Model candidates ranged from K=2 to K=10.",
    )
    add_paragraph(
        doc,
        "The current benchmark treats literal true K as a known generating condition, not as the sole target. "
        "Because noisy or overlapping generated states may not all be recoverable, the selected SPM-MS criterion "
        "is chosen by recovered-state F1: sensitivity rewards recovered true states and precision penalises extra "
        "estimated states. Squared K error is then used as the nonlinear K-count cost, so misses by several "
        "states are penalised more strongly than central-but-wrong estimates. By this recoverable-structure "
        f"target, {PRIMARY_SPM_CRITERION_TEXT} was tied for the best SPM-MS "
        f"criterion (F1 {fmt_pct(selected_criterion['F1'])}, sensitivity {fmt_pct(selected_criterion['sensitivity'])}, "
        f"precision {fmt_pct(selected_criterion['precision'])}, squared K error "
        f"{fmt_num(selected_criterion['sq_K_error'], 3)}). The literal exact-K winner was "
        f"{criterion_short_label(exact_winner['criterion'])}, but its exact-K gain was small "
        f"({fmt_pct(exact_winner['exact_K'])} versus {fmt_pct(selected_criterion['exact_K'])}) and came with "
        f"lower recovered-state F1 ({fmt_pct(exact_winner['F1'])}) and higher squared K error "
        f"({fmt_num(exact_winner['sq_K_error'], 3)}).",
    )

    add_heading(doc, "Descriptive report: criterion behaviour", 3)
    add_paragraph(
        doc,
        "The descriptive report shows why criterion choice should be stated explicitly. Raw criteria were strongly "
        "directional: K means with silhouette and raw SPM-MS free energy under-selected K, while raw covariance "
        "and GEV over-selected K. Free-energy elbow selected a mean K of 5.16 with a mean signed K difference "
        "of -0.34 states (estimated minus true), compared with K means mean selected K of 3.22 and signed "
        "difference of -2.28 states. Its squared K error was 1.88, compared with 6.55 for K means/silhouette. "
        "Its over-selection rate was 28.7%, but precision remained high because unmatched extra states were "
        "uncommon relative to the structure recovered. When F1 was plotted against "
        "true K, the selected SPM-MS criterion remained higher than K means across K=4-7.",
    )
    add_paragraph(
        doc,
        f"Criterion comparison is kept separate from the primary benchmark. Tables 1-2 and Figures 4-5 below "
        f"use only SPM-MS/{PRIMARY_SPM_CRITERION_TEXT} versus K-means/silhouette; they do not average SPM-MS "
        "across criteria.",
    )

    figures = [
        (
            criterion_plot,
            "Figure 2. Criterion-level selected K and signed K error. The dashed reference line marks the "
            "balanced simulation mean true K in the left panel and zero estimated-minus-true error in the right panel.",
        ),
        (
            SIM_FIG_DIR / "criterion_effects_with_ci.png",
            "Figure 3. Criterion-level simulation performance. Matrix-space LL, BIC, and ICL are reported "
            "alongside the legacy free-energy elbow and covariance criteria, "
            "whereas raw free energy underfit and raw covariance or GEV tended to overfit.",
        ),
        (
            f1_plot,
            "Figure 4. Recoverable-structure selected criterion per method: recovered-state F1 by true K. The plotted criteria "
            f"are silhouette for K means and {PRIMARY_SPM_CRITERION_TEXT} for SPM-MS; shaded bands show 95% "
            "confidence intervals.",
        ),
    ]
    for path, caption in figures:
        if path.exists():
            add_picture(doc, path)
            add_caption(doc, caption)

    doc.add_page_break()
    add_heading(doc, "Benchmark report: paired inference", 3)
    rows = []
    for montage in montages:
        exact = metric_row(benchmark, montage, "Exact K recovery")
        sq_err = metric_row(benchmark, montage, "Squared K error")
        f1 = metric_row(benchmark, montage, "Recovered-state F1")
        rows.append(
            [
                montage,
                exact["K means"],
                exact["SPM-MS"],
                sq_err["K means"],
                sq_err["SPM-MS"],
                f1["K means"],
                f1["SPM-MS"],
            ]
        )
    add_table(
        doc,
        [
            "Montage",
            f"Exact K {PRIMARY_KM_TABLE_LABEL}",
            f"Exact K {PRIMARY_SPM_TABLE_LABEL}",
            f"Sq. K error {PRIMARY_KM_TABLE_LABEL}",
            f"Sq. K error {PRIMARY_SPM_TABLE_LABEL}",
            f"F1 {PRIMARY_KM_TABLE_LABEL}",
            f"F1 {PRIMARY_SPM_TABLE_LABEL}",
        ],
        rows,
        widths=[1.2, 0.82, 0.82, 0.9, 0.9, 0.72, 0.72],
    )
    add_caption(doc, "Table 1. Primary paired simulation benchmark results by montage.")

    gains = []
    for montage in montages:
        exact = metric_row(benchmark, montage, "Exact K recovery")
        sq_err = metric_row(benchmark, montage, "Squared K error")
        f1 = metric_row(benchmark, montage, "Recovered-state F1")
        gains.append(
            f"{montage}: exact K {exact['Improvement']} ({exact['95% CI']}), "
            f"squared K error {sq_err['Improvement']} ({sq_err['95% CI']}), "
            f"F1 {f1['Improvement']} ({f1['95% CI']})"
        )
    add_paragraph(
        doc,
        "Paired block-bootstrap intervals and sign-flip tests supported the same conclusion: "
        + "; ".join(gains)
        + ". Holm-adjusted permutation p values were 0.002 for these primary contrasts.",
    )
    sens_gains = [
        metric_row(benchmark, montage, "Sensitivity")["Improvement"].replace("+", "").replace(" pp", "")
        for montage in montages
    ]
    add_paragraph(
        doc,
        "The gains came mainly from avoiding state collapse. SPM-MS increased recovered-state "
        f"sensitivity by {sens_gains[0]}, {sens_gains[1]}, and {sens_gains[2]} percentage points in full, "
        "19-channel, and 12-channel "
        "montages, respectively. Precision was lower for SPM-MS because it retained additional states "
        "that K means often omitted, but F1 and penalised map recovery still improved substantially.",
    )

    rows = []
    for montage in montages:
        identity = metric_row(benchmark, montage, "Template identity accuracy")
        backfit = metric_row(benchmark, montage, "Backfit label accuracy")
        runtime = metric_row(benchmark, montage, "Runtime")
        rows.append(
            [
                montage,
                identity["K means"],
                identity["SPM-MS"],
                backfit["K means"],
                backfit["SPM-MS"],
                runtime["K means"],
                runtime["SPM-MS"],
            ]
        )
    add_table(
        doc,
        [
            "Montage",
            f"Template ID {PRIMARY_KM_TABLE_LABEL}",
            f"Template ID {PRIMARY_SPM_TABLE_LABEL}",
            f"Backfit acc. {PRIMARY_KM_TABLE_LABEL}",
            f"Backfit acc. {PRIMARY_SPM_TABLE_LABEL}",
            f"Runtime {PRIMARY_KM_TABLE_LABEL}",
            f"Runtime {PRIMARY_SPM_TABLE_LABEL}",
        ],
        rows,
        widths=[1.2, 0.85, 0.85, 0.9, 0.9, 0.78, 0.78],
    )
    add_caption(doc, "Table 2. Template identity, mixture backfit, and runtime comparisons.")

    add_heading(doc, "Confusion and backfit diagnostics", 3)
    add_paragraph(
        doc,
        "The selected-criterion K-selection confusion plot is retained as the main confusion diagnostic. It "
        f"shows the important caveat behind literal exact-K results: SPM-MS/{PRIMARY_SPM_CRITERION_TEXT} often "
        "selected K=5 and still under-selected higher-K cases, but it avoided the stronger low-K collapse seen "
        "under K-means/silhouette while preserving high precision. This explains why recovered-state F1 and "
        "penalised map recovery are more meaningful primary endpoints than exact equality with the generating K.",
    )
    add_paragraph(
        doc,
        "The detailed backfit confusion matrices generated by the pipeline are treated as quality-control "
        "outputs rather than primary figures. Their off-diagonal structure depends on true-estimated cardinality "
        "mismatch, label matching, and empty-state padding, so reproducing every matrix in the manuscript would "
        "make a diagnostic look more inferential than it is. The manuscript therefore reports aggregate backfit "
        "label accuracy and label-weight error in Table 2, and uses the compact K-selection confusion plot to "
        "show the model-selection failure modes directly.",
    )
    figures = [
        (
            confusion_plot,
            "Figure 5. K-selection confusion by montage for the primary benchmark criteria only. Rows are true K "
            f"and columns are selected K, scaled within row. SPM-MS is {PRIMARY_SPM_CRITERION_TEXT} only; K-means is "
            "silhouette only.",
        ),
    ]
    for path, caption in figures:
        if path.exists():
            add_picture(doc, path)
            add_caption(doc, caption)


def add_lemon_results(doc, stats):
    clean_manifest = stats["clean_manifest"]
    fit = stats["fit"]
    model = stats["global_model"]
    hard_records = stats["hard_records"]
    hard_states = stats["hard_states"]
    common_channels = stats["common_channels"]
    pca = stats["pca"]
    pc_tests = stats["pc_tests"]
    demo = stats["demographics"]
    focused = stats["focused"]
    state_e = stats["state_e"]

    add_heading(doc, "LEMON hierarchical microstates", 2)
    add_paragraph(
        doc,
        f"The LEMON pipeline analysed {clean_manifest['participant'].nunique()} participants and "
        f"{len(clean_manifest)} clean resting EEG records "
        f"({(clean_manifest['condition'] == 'EC').sum()} eyes-closed, "
        f"{(clean_manifest['condition'] == 'EO').sum()} eyes-open) on "
        f"{len(common_channels)} common scalp channels. One manifest row with the literal condition label "
        "'condition' was excluded from condition-specific and behavioural summaries. The clean EC/EO "
        f"condition fits contained {int(fit.loc[fit['level'].eq('condition'), 'n_maps'].sum()):,} GFP-peak maps; "
        "the global SPM-MS fit used the configured 50,000-map pooled cap.",
    )

    selected = model.loc[model["selection_score"].idxmax()]
    add_paragraph(
        doc,
        f"Across K=4-7, the free-energy-covariance hierarchy criterion selected K={int(selected['K'])} "
        f"for the global model (selection score {fmt_num(selected['selection_score'], 3)}, "
        f"GEV {fmt_num(selected['gev'], 3)}, silhouette {fmt_num(selected['silhouette'], 3)}). "
        "This K was then propagated through condition, participant, and participant-condition fits using "
        "parent-template pseudo-priors and alignment to the EEG-Meta-Microstates topographies.",
    )

    rows = []
    for _, row in model.iterrows():
        rows.append(
            [
                int(row["K"]),
                f"{row['free_energy']:.1f}",
                fmt_num(row["silhouette"], 3),
                fmt_num(row["gev"], 3),
                f"{row['wss']:.1f}",
                fmt_num(row["selection_score"], 3),
            ]
        )
    add_table(
        doc,
        ["K", "Free energy", "Silhouette", "GEV", "WSS", "Selection score"],
        rows,
        widths=[0.45, 1.25, 0.9, 0.7, 1.1, 1.05],
    )
    add_caption(doc, "Table 3. Global LEMON model comparison.")

    global_row = fit[fit["level"].eq("global")].iloc[0]
    ec_row = fit[(fit["level"].eq("condition")) & (fit["condition"].eq("EC"))].iloc[0]
    eo_row = fit[(fit["level"].eq("condition")) & (fit["condition"].eq("EO"))].iloc[0]
    participant = fit[fit["level"].eq("participant")]
    pcond = fit[(fit["level"].eq("participant_condition")) & (fit["condition"].isin(["EC", "EO"]))]
    rows = [
        [
            "Global",
            f"{int(global_row['n_units'])} records",
            f"{int(global_row['n_maps']):,}",
            int(global_row["K_estimated"]),
            fmt_num(global_row["template_mean_corr"], 3),
            int(global_row["template_strong_matches"]),
        ],
        [
            "Eyes closed",
            f"{int(ec_row['n_units'])} records",
            f"{int(ec_row['n_maps']):,}",
            int(ec_row["K_estimated"]),
            fmt_num(ec_row["template_mean_corr"], 3),
            int(ec_row["template_strong_matches"]),
        ],
        [
            "Eyes open",
            f"{int(eo_row['n_units'])} records",
            f"{int(eo_row['n_maps']):,}",
            int(eo_row["K_estimated"]),
            fmt_num(eo_row["template_mean_corr"], 3),
            int(eo_row["template_strong_matches"]),
        ],
        [
            "Participant",
            f"{len(participant)} participants",
            f"median {participant['n_maps'].median():.0f}",
            "all 6",
            f"mean {participant['template_mean_corr'].mean():.3f}",
            f"median {participant['template_strong_matches'].median():.0f}",
        ],
        [
            "Participant-condition",
            f"{len(pcond)} records",
            f"median {pcond['n_maps'].median():.0f}",
            "all 6",
            f"mean {pcond['template_mean_corr'].mean():.3f}",
            f"median {pcond['template_strong_matches'].median():.0f}",
        ],
    ]
    add_table(
        doc,
        ["Level", "Units", "GFP maps", "K", "Template corr.", "Strong matches"],
        rows,
        widths=[1.35, 1.2, 1.05, 0.55, 0.95, 0.85],
    )
    add_caption(doc, "Table 4. Hierarchical SPM-MS fit summary.")

    panel = build_lemon_topography_panel()
    add_picture(doc, panel)
    add_caption(
        doc,
        "Figure 6. LEMON microstate centres from the current hierarchy. The global solution selected "
        "six maps; eyes-closed and eyes-open condition fits preserved the same K and remained strongly "
        "aligned with EEG-Meta-Microstates templates.",
    )

    doc.add_page_break()
    add_heading(doc, "LEMON backfit dynamics and behavioural associations", 3)
    add_paragraph(
        doc,
        f"Current LEMON backfitting used hard template assignment for all {len(hard_records)} clean EC/EO records. "
        "Gaussian mixture backfit entries are present in the output tables but unavailable in this run, so they "
        "are not interpreted. Hard-backfit record duration averaged "
        f"{hard_records['duration_s'].mean():.1f} s (SD {hard_records['duration_s'].std():.1f}; "
        f"range {hard_records['duration_s'].min():.1f}-{hard_records['duration_s'].max():.1f}).",
    )

    state_summary = (
        hard_states.groupby(["condition", "template_label"], observed=True)
        .agg(
            n=("occupancy", "size"),
            occupancy_mean=("occupancy", "mean"),
            occupancy_sd=("occupancy", "std"),
            occurrence_rate=("occurrence_rate_hz", "mean"),
            template_corr=("template_match_abs_correlation", "mean"),
        )
        .reset_index()
        .sort_values(["condition", "template_label"])
    )
    rows = []
    for _, row in state_summary.iterrows():
        rows.append(
            [
                row["condition"],
                row["template_label"],
                int(row["n"]),
                f"{row['occupancy_mean']:.3f} ({row['occupancy_sd']:.3f})",
                fmt_num(row["occurrence_rate"], 2),
                fmt_num(row["template_corr"], 3),
            ]
        )
    add_table(
        doc,
        ["Condition", "Label", "Records", "Occupancy mean (SD)", "Rate Hz", "Template corr."],
        rows,
        widths=[0.85, 0.55, 0.7, 1.35, 0.75, 1.0],
    )
    add_caption(doc, "Table 5. Hard-backfit LEMON state dynamics by aligned template label.")

    pc_scores = pd.read_csv(BEHAV_DIR / "psychometric_pc_scores.csv")
    add_paragraph(
        doc,
        f"Psychometric PCA retained {len(pca)} components to explain "
        f"{100 * pca['cumulative_variance_ratio'].iloc[-1]:.1f}% of behavioural-score variance "
        f"across {len(pc_scores)} participants with retained behavioural scores. PC-microstate screening "
        f"ran {len(pc_tests)} tests; the smallest uncorrected p value was {fmt_p(pc_tests['p_value'].min())}, "
        f"but the smallest FDR q value was {fmt_num(pc_tests['fdr_q_value'].min(), 3)}, so no PC association "
        "survived correction.",
    )

    top_demo = demo.sort_values("fdr_q_value").head(3)
    top_focused = focused.sort_values("p_value").head(1).iloc[0]
    yfas = state_e[
        (state_e["behaviour"].str.contains("yfas", case=False, na=False))
        & (state_e["microstate_metric"] == "state__EO__E__occupancy")
    ]
    yfas_text = ""
    if not yfas.empty:
        yr = yfas.iloc[0]
        yfas_text = (
            f"YFAS symptom count versus EO state E occupancy was only an uncorrected trend "
            f"(Spearman rho {yr['spearman_rho']:.3f}, p={fmt_p(yr['spearman_p'])})."
        )
    rows = [
        [
            "Psychometric PCs",
            f"{len(pc_tests)} tests",
            "min q=" + fmt_num(pc_tests["fdr_q_value"].min(), 3),
            "No FDR-significant association",
        ]
    ]
    for _, row in top_demo.iterrows():
        rows.append(
            [
                row["feature_name"].replace("state__", "").replace("__", " "),
                f"n={int(row['n'])}",
                f"rho={row['effect_size']:.3f}, p={fmt_p(row['p_value'])}, q={fmt_p(row['fdr_q_value'])}",
                "Age association survived FDR",
            ]
        )
    rows.append(
        [
            top_focused["behaviour_construct"] + " / " + top_focused["microstate_feature"].replace("state__", ""),
            f"n={int(top_focused['n'])}",
            f"rho={top_focused['spearman_rho']:.3f}, p={fmt_p(top_focused['p_value'])}, q={fmt_p(top_focused['fdr_q_value'])}",
            "Exploratory only",
        ]
    )
    if yfas_text:
        rows.append(["State E focused plot", "", yfas_text, "Not treated as robust evidence"])
    add_table(
        doc,
        ["Analysis", "N/tests", "Effect", "Interpretation"],
        rows,
        widths=[1.85, 0.8, 2.1, 1.35],
    )
    add_caption(doc, "Table 6. Behavioural and demographic analyses from the current hard-backfit features.")


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    stats = load_stats()

    doc = Document()
    style_doc(doc)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Variational Bayesian Gaussian Mixture Microstates for EEG")
    p.add_run("\n")
    p.add_run("Simulation Recovery and Hierarchical LEMON Validation")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Rohan Kandasamy")
    doc.add_paragraph()

    add_heading(doc, "Abstract", 1)
    add_bold_lead(
        doc,
        "Background. ",
        "EEG microstate analysis usually combines clustering of global-field-power peaks with a separate "
        "choice of state count. This can bias downstream temporal metrics when weak or overlapping states "
        "are omitted.",
    )
    add_bold_lead(
        doc,
        "Methods. ",
        "The current SPM-MS pipeline fits scalp-map mixtures with the SPM variational Bayes Gaussian mixture "
        "implementation, evaluates candidate K values, aligns fitted maps to EEG-Meta-Microstates templates, "
        "and supports parent-seeded hierarchical fits. Performance was tested in a ground-truth simulation "
        "and then applied to the LEMON resting EEG dataset.",
    )
    add_bold_lead(
        doc,
        "Results. ",
        "Across 1,680 matched simulation blocks, SPM-MS is compared with K means using paired montage-specific "
        "exact-K recovery, squared K error, recovered-state F1, and map-recovery diagnostics. In LEMON, the hierarchy selected K=6, fitted 203 participants and "
        "405 clean eyes-closed/eyes-open records on 59 common channels, and produced stable template-aligned "
        "condition and participant-condition maps. Behavioural PC screening did not survive FDR correction, "
        "whereas three eyes-open occurrence-rate features showed FDR-significant age associations.",
    )
    add_bold_lead(
        doc,
        "Conclusion. ",
        "SPM-MS improves ground-truth state recovery mainly by preserving states that K means tends to collapse. "
        "The current real-data hierarchy is usable for descriptive LEMON microstate dynamics, but psychometric "
        "associations should remain exploratory until replicated with corrected endpoints.",
    )

    add_heading(doc, "Introduction", 1)
    add_paragraph(
        doc,
        "Resting EEG microstates summarise the scalp potential field as a sequence of short-lived, quasi-stable "
        "topographies. Their appeal is practical: a high-dimensional time series is reduced to a small set of "
        "state maps and temporal descriptors such as occupancy, occurrence rate, duration, transition structure, "
        "and map fit. The difficult part is not only estimating maps, but deciding how many maps should exist "
        "and whether the resulting states align across recordings.",
    )
    add_paragraph(
        doc,
        "The current pipeline, SPM-MS, treats microstate extraction as a variational Bayesian Gaussian mixture "
        "problem over normalised scalp maps. Candidate models are compared using matrix-space free energy, "
        "LL, BIC, ICL, covariance, silhouette, and combined elbow criteria, then aligned to the "
        "EEG-Meta-Microstates templates. This "
        "manuscript reports the pipeline as it currently runs: first on simulated EEG with known microstate "
        "ground truth, and then on the LEMON resting EEG dataset using a hierarchical global-to-condition-to-"
        "participant fitting strategy.",
    )

    add_heading(doc, "Methods", 1)
    add_heading(doc, "SPM-MS fitting", 2)
    add_paragraph(
        doc,
        "SPM-MS extracts GFP peaks, optionally filters and smooths the scalp maps, and fits candidate Gaussian "
        "mixture models using the SPM variational Bayes mixture implementation. Maps are treated as polarity-"
        "invariant spatial patterns by matching on absolute spatial correlation after normalisation. Candidate "
        "K values are scored after refinement in polarity-invariant matrix space using free energy, LL, BIC, "
        f"and ICL. The simulation benchmark uses the {PRIMARY_SPM_CRITERION_TEXT} criterion for the paired "
        "SPM-MS comparison; literal exact-K recovery is reported as a diagnostic. "
        "The LEMON hierarchy uses the configured free-energy-covariance criterion.",
    )
    add_paragraph(
        doc,
        "Template alignment is performed against EEG-Meta-Microstates reference maps. For hierarchical LEMON "
        "fits, parent solutions seed child solutions by pseudo-count priors: the global fit seeds condition and "
        "participant fits, and participant/condition templates seed participant-condition fits. This keeps labels "
        "comparable across the hierarchy without requiring every record to express every template label.",
    )

    add_heading(doc, "Ground-truth simulation", 2)
    add_paragraph(
        doc,
        "Simulated EEG was generated from known template states, with true K values 4, 5, 6, and 7; SNR values "
        "-6, -3, -1, 0, 1, 3, and 6 dB; 10 replicates; and overlap probabilities 0 and 0.5. Overlap events lasted "
        "10-40 ms and used strength 0.5. Each generated full-montage EEG block was projected to full 71-channel, "
        "19-channel 10-20-20, and 12-channel 10-20-12 montages. Both SPM-MS and Koenig-style K means were fit to "
        "the same data with K candidates 2-10.",
    )
    demo_panel = build_pipeline_demonstrator_panel()
    add_picture(doc, demo_panel)
    add_caption(
        doc,
        "Figure 1. Pipeline demonstrator. A four-channel excerpt from a regenerated simulated EEG block "
        "shows the noisy scalp signal and the underlying true microstate sequence; the flow chart summarises "
        "the simulation-to-report path used by the benchmark.",
    )
    add_paragraph(
        doc,
        "Primary recovery metrics were exact K recovery, squared K error, matched and penalised map recovery, "
        "state sensitivity, precision, F1, template identity accuracy, mixture-backfit label accuracy, label-weight "
        "mean absolute error, overlap-pair accuracy, and runtime. Statistical inference used the recent benchmark "
        "report: within-montage block-bootstrap 95% confidence intervals with 2,000 resamples and one-sided paired "
        "sign-flip permutation tests with 20,000 sign flips, Holm-adjusted across reported simulation metrics.",
    )

    add_heading(doc, "LEMON hierarchy", 2)
    add_paragraph(
        doc,
        "LEMON EEG data were read from preprocessed EEGLAB .set files. The pipeline used average referencing, "
        "2-20 Hz bandpass filtering, common-channel harmonisation, interpolation of missing channels, spatial "
        "smoothing, GFP peak extraction, and GFP outlier rejection. PO9 and PO10 were excluded from the common "
        "channel set. The final clean EC/EO manifest contained 203 participants and 405 records on 59 common "
        "scalp channels.",
    )
    add_paragraph(
        doc,
        "A pooled global model was fit with a 50,000-map cap across candidate K values 4-7. Condition, participant, "
        "and participant-condition models inherited the selected global K and parent-template priors. Participant-"
        "condition maps were then hard-backfit to the full EEG records to compute occupancy, percentage-present, "
        "mean GFP, occurrence count, occurrence rate, entropy, and pairwise transition information. Gaussian mixture "
        "backfit placeholders are present in the output tables but unavailable for the current LEMON run.",
    )

    add_heading(doc, "Behavioural statistics", 2)
    add_paragraph(
        doc,
        "Psychometric scores were reduced by PCA, retaining components up to approximately 80% cumulative explained "
        "variance. Backfit features were screened against psychometric PCs using Pearson correlations and against "
        "demographic variables using the analysis type specified in the output tables, with FDR correction. Focused "
        "interoception and dissociation analyses used Spearman and partial Spearman correlations adjusted for age "
        "and gender where available. These focused analyses are treated as exploratory unless they survive FDR "
        "correction.",
    )

    add_heading(doc, "Results", 1)
    add_simulation_results(doc, stats)
    add_lemon_results(doc, stats)

    add_heading(doc, "Discussion", 1)
    add_paragraph(
        doc,
        "The simulation experiment supports the current SPM-MS pipeline as a better recovery tool than K means "
        "with silhouette selection when the ground-truth state count is known. The most important gain is not that "
        "exact K recovery becomes perfect; it remains around one quarter of blocks. The gain is that squared K "
        "error drops substantially and recovered-state F1 rises to about 90% across all montage projections. This is a "
        "more useful target for downstream EEG analysis because temporal and behavioural features depend on whether "
        "states are retained and aligned, not only on exact K equality.",
    )
    add_paragraph(
        doc,
        "The LEMON hierarchy shows that the same pipeline can be run on real resting EEG at cohort scale. The "
        "global K=6 solution aligned strongly with reference templates and propagated through condition and "
        "participant-condition fits without changing K. Eyes-closed and eyes-open state dynamics were descriptive "
        "rather than confirmatory: state labels were stable enough to summarise occupancy and occurrence, but "
        "behavioural PC screening did not produce corrected psychometric associations. The robust behavioural "
        "signal in the current output is demographic, with eyes-open B, F, and G occurrence rates increasing with "
        "age midpoint after FDR correction.",
    )
    add_paragraph(
        doc,
        "Several limitations matter. First, the simulation generator defines a specific family of state overlap, "
        "noise, and template mixing; other artefact models could change the relative difficulty of the task. "
        "Second, exact K recovery remains modest, so model-selection outputs should be interpreted with uncertainty "
        "rather than as absolute truth. Third, the current LEMON run lacks usable Gaussian mixture backfit outputs, "
        "so real-data temporal metrics are based on hard state assignment. Finally, the psychometric analyses are "
        "high-dimensional and largely exploratory; the corrected null results are as important as the uncorrected "
        "trends.",
    )
    add_heading(doc, "Conclusion", 1)
    add_paragraph(
        doc,
        "For the current pipeline, the strongest evidence is methodological: SPM-MS improves simulated state recovery "
        "and remains robust under reduced montage density. The LEMON hierarchy provides a coherent real-data "
        "demonstration with K=6 template-aligned microstates across 203 participants, but behavioural interpretation "
        "should focus on the corrected age associations and avoid overclaiming exploratory psychometric trends.",
    )

    refs = read_bibliography()
    if refs:
        add_heading(doc, "References", 1)
        for ref in refs:
            add_paragraph(doc, ref)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
